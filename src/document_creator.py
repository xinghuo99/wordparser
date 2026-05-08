from docx import Document
from docx.shared import Pt


def create_test_document(output_path: str):
    doc = Document()
    
    section = doc.sections[0]
    header = section.header
    header_paragraph = header.paragraphs[0]
    header_paragraph.text = "文档标题 - Word Parser 测试文档"
    header_paragraph.style = doc.styles['Header']
    
    doc.add_heading('第一章 概述', level=1)
    doc.add_paragraph('这是第一章的正文内容，介绍文档的整体概述。')
    doc.add_paragraph('继续添加更多内容到第一章。')
    
    doc.add_heading('1.1 项目背景', level=2)
    doc.add_paragraph('本节介绍项目的背景信息，说明为什么需要这个项目。')
    
    table1 = doc.add_table(rows=3, cols=3)
    table1.style = 'Table Grid'
    hdr_cells = table1.rows[0].cells
    hdr_cells[0].text = '序号'
    hdr_cells[1].text = '项目名称'
    hdr_cells[2].text = '状态'
    row_cells = table1.rows[1].cells
    row_cells[0].text = '1'
    row_cells[1].text = '项目A'
    row_cells[2].text = '进行中'
    row_cells = table1.rows[2].cells
    row_cells[0].text = '2'
    row_cells[1].text = '项目B'
    row_cells[2].text = '已完成'
    
    doc.add_heading('1.2 项目目标', level=2)
    doc.add_paragraph('本节描述项目的主要目标和预期成果。')
    
    doc.add_heading('第二章 技术方案', level=1)
    doc.add_paragraph('本章详细描述技术方案的设计和实现细节。')
    
    doc.add_heading('2.1 技术选型', level=2)
    doc.add_paragraph('选择合适的技术栈对于项目成功至关重要。')
    
    doc.add_heading('2.1.1 前端技术', level=3)
    doc.add_paragraph('前端采用现代Web技术栈，包括React和TypeScript。')
    
    table2 = doc.add_table(rows=2, cols=2)
    table2.style = 'Table Grid'
    table2.rows[0].cells[0].text = '技术'
    table2.rows[0].cells[1].text = '版本'
    table2.rows[1].cells[0].text = 'React'
    table2.rows[1].cells[1].text = '18.x'
    
    doc.add_heading('2.1.2 后端技术', level=3)
    doc.add_paragraph('后端采用Python和FastAPI构建RESTful API。')
    
    doc.add_heading('2.2 架构设计', level=2)
    doc.add_paragraph('系统采用微服务架构，支持高并发和可扩展性。')
    
    doc.add_heading('第三章 修订记录', level=1)
    doc.add_paragraph('本文档的修订历史记录如下：')
    
    revision_table = doc.add_table(rows=4, cols=4)
    revision_table.style = 'Table Grid'
    hdr_cells = revision_table.rows[0].cells
    hdr_cells[0].text = '版本号'
    hdr_cells[1].text = '修订日期'
    hdr_cells[2].text = '修订人'
    hdr_cells[3].text = '修订内容'
    row_cells = revision_table.rows[1].cells
    row_cells[0].text = 'V1.0'
    row_cells[1].text = '2024-01-15'
    row_cells[2].text = '张三'
    row_cells[3].text = '初始版本'
    row_cells = revision_table.rows[2].cells
    row_cells[0].text = 'V1.1'
    row_cells[1].text = '2024-02-20'
    row_cells[2].text = '李四'
    row_cells[3].text = '添加技术选型章节'
    row_cells = revision_table.rows[3].cells
    row_cells[0].text = 'V1.2'
    row_cells[2].text = '王五'
    row_cells[3].text = '更新架构设计'
    
    doc.add_heading('第四章 附录', level=1)
    doc.add_paragraph('本章节包含文档的附录信息。')
    
    doc.add_paragraph('修订记录')
    doc.add_paragraph('以下是文档的详细修订说明：')
    
    appendix_table = doc.add_table(rows=2, cols=2)
    appendix_table.style = 'Table Grid'
    appendix_table.rows[0].cells[0].text = '版本'
    appendix_table.rows[0].cells[1].text = '说明'
    appendix_table.rows[1].cells[0].text = 'V1.3'
    appendix_table.rows[1].cells[1].text = '添加附录内容'
    
    doc.save(output_path)
    print(f"测试文档已创建: {output_path}")


if __name__ == '__main__':
    create_test_document('test_document.docx')