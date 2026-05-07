from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph
from typing import List, Dict, Union, Optional

class Section:
    """章节数据结构"""
    def __init__(self, title: str, level: int):
        self.title = title.strip()
        self.level = level
        self.content = []          # 混合存储：str(段落文本) 或 list(表格二维列表)
        self.tables = []           # 单独存放所有表格的二维列表，便于查找

def extract_sections(docx_path: str) -> List[Section]:
    """
    从 Word 文档提取章节。
    返回 Section 列表，按文档顺序排列。
    文档开头的非标题内容会归入 title 为空的特殊章节。
    """
    doc = Document(docx_path)
    sections = []
    current_section = Section(title="", level=0)   # 初始章节，存放第一个标题之前的零散内容

    # 遍历文档 body 中的所有元素（段落和表格）
    for child in doc.element.body:
        if child.tag == qn('w:p'):          # 段落
            para = Paragraph(child, doc)
            style_name = para.style.name if para.style else ''
            # 判断是否为标题（以 Heading 开头）
            if style_name.startswith('Heading'):
                # 提取标题级别（如 Heading 1 -> 1）
                try:
                    level = int(style_name.split()[-1])
                except ValueError:
                    level = 0
                # 保存当前章节（若标题非空）
                if current_section.title:
                    sections.append(current_section)
                # 开始新章节
                current_section = Section(title=para.text, level=level)
                # 标题本身不加入内容，仅作为章节名
            else:
                # 普通段落：将文本加入当前章节内容
                text = para.text.strip()
                if text:  # 忽略空段落
                    current_section.content.append(text)
        elif child.tag == qn('w:tbl'):      # 表格
            table = Table(child, doc)
            # 解析表格为二维列表
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            # 同时存入 content 和 tables 列表
            current_section.content.append(table_data)
            current_section.tables.append(table_data)

    # 添加最后一个章节
    if current_section:
        sections.append(current_section)

    # 可选：去掉起始的空章节（若文档开头确实无标题且无内容）
    if sections and sections[0].title == "" and not sections[0].content:
        sections.pop(0)
    return sections

def find_section_by_title(sections: List[Section], title: str) -> Optional[Section]:
    """
    按章节标题查找，返回第一个匹配的 Section 对象。
    参数 title 将去除首尾空格后与 Section.title 进行精确匹配。
    """
    clean_title = title.strip()
    for sec in sections:
        if sec.title == clean_title:
            return sec
    return None

# ---------- 使用示例 ----------
if __name__ == "__main__":
    # 提取所有章节
    sections = extract_sections("example.docx")
    print(f"共提取 {len(sections)} 个章节：")
    for sec in sections:
        print(f"  [{sec.level}] {sec.title}  内容段落数: {sum(1 for c in sec.content if isinstance(c, str))}  表格数: {len(sec.tables)}")

    # 按名称查找某个章节
    target = find_section_by_title(sections, "第三章 研究方法")
    if target:
        print(f"\n找到章节「{target.title}」，内容如下：")
        for item in target.content:
            if isinstance(item, str):
                print("  段落：", item[:60])       # 仅显示前60字符
            else:  # 表格
                print("  表格：", item[0])        # 显示表头行
    else:
        print("未找到该章节")