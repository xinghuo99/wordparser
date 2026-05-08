from docx import Document
from docx.shared import Pt
from typing import List, Dict, Any, Optional


class Section:
    def __init__(self, title: str, level: int, content: str = "", tables: List[Dict] = None):
        self.title = title
        self.level = level
        self.content = content
        self.tables = tables if tables else []
        self.children = []

    def add_child(self, child: 'Section'):
        self.children.append(child)

    def to_dict(self) -> Dict[str, Any]:
        return {
            'title': self.title,
            'level': self.level,
            'content': self.content,
            'tables': self.tables,
            'children': [child.to_dict() for child in self.children]
        }


class WordParser:
    def __init__(self, doc_path: str):
        self.doc_path = doc_path
        self.document = None
        self.sections = []
        self.headers = []
        self.style_level_map = {
            'Heading 1': 1,
            'Heading 2': 2,
            'Heading 3': 3,
            'Heading 4': 4,
            'Heading 5': 5,
            'Heading 6': 6,
        }

    def load_document(self):
        self.document = Document(self.doc_path)

    def _get_paragraph_level(self, paragraph) -> Optional[int]:
        style_name = paragraph.style.name
        return self.style_level_map.get(style_name)

    def _parse_table(self, table) -> List[List[str]]:
        table_data = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                row_data.append(cell.text.strip())
            table_data.append(row_data)
        return table_data

    def extract_sections(self) -> List[Section]:
        if not self.document:
            self.load_document()

        self.sections = []
        stack = []
        current_section = None
        
        all_elements = []
        for paragraph in self.document.paragraphs:
            all_elements.append(('paragraph', paragraph))
        
        for table in self.document.tables:
            all_elements.append(('table', table))
        
        all_elements.sort(key=lambda x: self._get_element_position(x[1]))

        for elem_type, element in all_elements:
            if elem_type == 'paragraph':
                paragraph = element
                level = self._get_paragraph_level(paragraph)
                
                if level is not None:
                    new_section = Section(title=paragraph.text.strip(), level=level)
                    
                    while stack and stack[-1].level >= level:
                        stack.pop()
                    
                    if stack:
                        stack[-1].add_child(new_section)
                    else:
                        self.sections.append(new_section)
                    
                    stack.append(new_section)
                    current_section = new_section
                elif current_section is not None:
                    current_section.content += paragraph.text.strip() + '\n'
            elif elem_type == 'table':
                if current_section is not None:
                    table_data = self._parse_table(element)
                    current_section.tables.append(table_data)

        for section in self.sections:
            section.content = section.content.strip()

        return self.sections

    def _get_element_position(self, element):
        if hasattr(element, '_element'):
            elem = element._element
        else:
            elem = element
        
        for i, child in enumerate(self.document.element.body):
            if child is elem:
                return i
        return float('inf')

    def extract_headers(self) -> List[Dict[str, Any]]:
        """提取文档中所有页眉内容，支持多种页眉格式"""
        if not self.document:
            self.load_document()

        self.headers = []
        
        for section_idx, section in enumerate(self.document.sections):
            headers_info = {
                'section_index': section_idx,
                'headers': []
            }
            
            header_types = [
                ('default', section.header),
                ('first_page', section.first_page_header),
                ('even_page', section.even_page_header)
            ]
            
            for header_type, header in header_types:
                if header is not None:
                    content = []
                    tables = []
                    
                    for paragraph in header.paragraphs:
                        text = paragraph.text.strip()
                        if text:
                            content.append(text)
                    
                    for table in header.tables:
                        tables.append(self._parse_table(table))
                    
                    if content or tables:
                        headers_info['headers'].append({
                            'type': header_type,
                            'content': '\n'.join(content),
                            'tables': tables
                        })
            
            if headers_info['headers']:
                self.headers.append(headers_info)
        
        return self.headers

    def get_all_headers_text(self) -> str:
        """获取所有页眉的文本内容，合并为一个字符串"""
        if not self.headers:
            self.extract_headers()
        
        all_text = []
        for section_info in self.headers:
            for header in section_info['headers']:
                if header['content']:
                    all_text.append(header['content'])
        
        return '\n'.join(all_text)

    def find_section_by_name(self, name: str, sections: Optional[List[Section]] = None) -> Optional[Section]:
        if sections is None:
            sections = self.sections

        for section in sections:
            if section.title == name:
                return section
            
            if section.children:
                found = self.find_section_by_name(name, section.children)
                if found:
                    return found
        
        return None

    def get_section_content(self, section: Section) -> Dict[str, Any]:
        return {
            'title': section.title,
            'level': section.level,
            'content': section.content,
            'tables': section.tables,
            'children': [self.get_section_content(child) for child in section.children]
        }

    def _table_to_text(self, table: List[List[str]]) -> str:
        """将表格内容转换为文字描述，按照表头+值的形式"""
        if not table or len(table) < 2:
            return ""
        
        headers = table[0]
        result_lines = []
        
        for row in table[1:]:
            row_desc = []
            for i, header in enumerate(headers):
                if i < len(row):
                    value = row[i]
                    if value:
                        row_desc.append(f"{header}: {value}")
            if row_desc:
                result_lines.append("；".join(row_desc))
        
        return "。\n".join(result_lines) + "。" if result_lines else ""

    def get_full_document_text(self, include_headers: bool = True) -> str:
        """获取整个文档的内容，文字加表格（表格内容组装成文字描述）"""
        if not self.sections:
            self.extract_sections()
        
        parts = []
        
        if include_headers:
            headers_text = self.get_all_headers_text()
            if headers_text:
                parts.append(headers_text)
        
        def process_section(section: Section, indent: int = 0):
            prefix = "    " * indent
            title_line = f"{prefix}{section.title}"
            parts.append(title_line)
            
            if section.content:
                content_lines = section.content.split('\n')
                for line in content_lines:
                    if line.strip():
                        parts.append(f"{prefix}    {line.strip()}")
            
            for table in section.tables:
                table_text = self._table_to_text(table)
                if table_text:
                    table_lines = table_text.split('\n')
                    for line in table_lines:
                        parts.append(f"{prefix}    表格内容：{line}")
            
            for child in section.children:
                process_section(child, indent + 1)
        
        for section in self.sections:
            process_section(section)
        
        return '\n'.join(parts)

    def get_section_text_by_name(self, name: str, include_children: bool = True) -> str:
        """按章节名称获取章节内容，文字加表格（表格内容组装成文字描述）"""
        if not self.sections:
            self.extract_sections()
        
        section = self.find_section_by_name(name)
        if not section:
            return f"未找到章节: {name}"
        
        parts = []
        
        def process_section(sec: Section, indent: int = 0):
            prefix = "    " * indent
            title_line = f"{prefix}{sec.title}"
            parts.append(title_line)
            
            if sec.content:
                content_lines = sec.content.split('\n')
                for line in content_lines:
                    if line.strip():
                        parts.append(f"{prefix}    {line.strip()}")
            
            for table in sec.tables:
                table_text = self._table_to_text(table)
                if table_text:
                    table_lines = table_text.split('\n')
                    for line in table_lines:
                        parts.append(f"{prefix}    表格内容：{line}")
            
            if include_children:
                for child in sec.children:
                    process_section(child, indent + 1)
        
        process_section(section)
        return '\n'.join(parts)

    def get_revision_records(self) -> tuple:
        """
        获取修订记录的内容，包括后面紧跟的文字和表格（表格内容组装成文字）
        
        支持两种情况：
        1. 修订记录是章节标题（Heading样式）
        2. 修订记录是普通文字段落
        
        :return: 元组 (content, has_table)
                 content: 修订记录内容字符串
                 has_table: 如果修订记录紧跟表格则为True，否则为False
        """
        if not self.sections:
            self.extract_sections()
        
        revision_section = None
        for section in self.sections:
            if '修订记录' in section.title:
                revision_section = section
                break
            if section.children:
                for child in section.children:
                    if '修订记录' in child.title:
                        revision_section = child
                        break
        
        if revision_section:
            return self._extract_section_revision_content(revision_section)
        
        return self._extract_paragraph_revision_content()

    def _extract_section_revision_content(self, section) -> tuple:
        """从章节中提取修订记录内容"""
        has_table = len(section.tables) > 0
        parts = []
        parts.append(section.title)
        
        if section.content:
            content_lines = section.content.split('\n')
            for line in content_lines:
                if line.strip():
                    parts.append(f"    {line.strip()}")
        
        for table in section.tables:
            table_text = self._table_to_text(table)
            if table_text:
                table_lines = table_text.split('\n')
                for line in table_lines:
                    parts.append(f"    表格内容：{line}")
        
        return '\n'.join(parts), has_table

    def _extract_paragraph_revision_content(self) -> tuple:
        """从普通文字段落中提取修订记录内容"""
        if not self.document:
            self.load_document()
        
        all_elements = []
        for paragraph in self.document.paragraphs:
            all_elements.append(('paragraph', paragraph))
        
        for table in self.document.tables:
            all_elements.append(('table', table))
        
        all_elements.sort(key=lambda x: self._get_element_position(x[1]))
        
        revision_start_idx = -1
        for i, (elem_type, element) in enumerate(all_elements):
            if elem_type == 'paragraph':
                if '修订记录' in element.text:
                    revision_start_idx = i
                    break
        
        if revision_start_idx == -1:
            return "", False
        
        parts = []
        has_table = False
        
        for i in range(revision_start_idx, len(all_elements)):
            elem_type, element = all_elements[i]
            
            if elem_type == 'paragraph':
                text = element.text.strip()
                if not text:
                    continue
                
                if _is_new_section(element) and i > revision_start_idx:
                    break
                
                parts.append(text)
            elif elem_type == 'table':
                has_table = True
                table_data = self._parse_table(element)
                table_text = self._table_to_text(table_data)
                if table_text:
                    table_lines = table_text.split('\n')
                    for line in table_lines:
                        parts.append(f"    表格内容：{line}")
        
        return '\n'.join(parts), has_table


    def extract_sections_to_list(self) -> List[Dict[str, str]]:
        """
        按照章节提取文档内容，返回列表格式
        每个列表元素是一个字典，包含：
        - 'title': 章节名称
        - 'content': 章节内容（包含文字和表格转文字，按文档顺序）
        
        如果包含子章节，content字段包含所有子章节和孙子章节的内容
        """
        if not self.sections:
            self.extract_sections()
        
        result = []
        
        def process_section(section: Section, indent: int = 0):
            content_parts = []
            
            if section.content:
                content_lines = section.content.split('\n')
                for line in content_lines:
                    if line.strip():
                        content_parts.append(line.strip())
            
            for table in section.tables:
                table_text = self._table_to_text(table)
                if table_text:
                    content_parts.append(table_text)
            
            for child in section.children:
                child_content = self._get_section_full_content(child)
                content_parts.append(child_content)
            
            section_dict = {
                'title': section.title,
                'content': '\n'.join(content_parts) if content_parts else ''
            }
            
            result.append(section_dict)
            
            for child in section.children:
                process_section(child, indent + 1)
        
        for section in self.sections:
            process_section(section)
        
        return result

    def _get_section_full_content(self, section: Section) -> str:
        """获取章节的完整内容，包括标题、文字、表格和所有子章节"""
        parts = []
        parts.append(section.title)
        
        if section.content:
            content_lines = section.content.split('\n')
            for line in content_lines:
                if line.strip():
                    parts.append(line.strip())
        
        for table in section.tables:
            table_text = self._table_to_text(table)
            if table_text:
                parts.append(table_text)
        
        for child in section.children:
            parts.append(self._get_section_full_content(child))
        
        return '\n'.join(parts)

    @staticmethod
    def find_section_content(sections_list: List[Dict[str, str]], section_name: str) -> str:
        """
        根据章节名称在章节列表中查找对应的章节内容
        
        :param sections_list: 章节列表，由 extract_sections_to_list() 返回
        :param section_name: 要查找的章节名称
        :return: 找到的章节内容，如果未找到返回空字符串
        """
        if not sections_list or not isinstance(section_name, str) or not section_name:
            return ""
        
        for section in sections_list:
            if isinstance(section, dict) and 'title' in section and section['title'] == section_name:
                return section.get('content', "")
        
        return ""


def _is_new_section(paragraph) -> bool:
    """判断段落是否是新章节的开始（Heading样式）"""
    style_name = paragraph.style.name
    return style_name.startswith('Heading')
