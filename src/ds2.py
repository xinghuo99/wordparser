import re
from docx import Document
from typing import List, Dict, Any, Union


class ChapterNode:
    """章节节点，保存标题、级别、内容和子章节"""
    def __init__(self, title: str = "", level: int = 0):
        self.title = title
        self.level = level
        self.paragraphs: List[str] = []   # 该章节下的段落文本
        self.tables: List[List[List[str]]] = []  # 表格内容，每个表格是二维列表
        self.children: List[ChapterNode] = []

    def add_paragraph(self, text: str):
        """添加一条段落文本（自动过滤空字符串）"""
        if text.strip():
            self.paragraphs.append(text.strip())

    def add_table(self, table_data: List[List[str]]):
        """添加一个表格的二维数据"""
        self.tables.append(table_data)

    def to_dict(self) -> Dict[str, Any]:
        """递归转为字典，便于 JSON 输出"""
        return {
            "level": self.level,
            "title": self.title,
            "paragraphs": self.paragraphs,
            "tables": self.tables,
            "children": [child.to_dict() for child in self.children]
        }


def extract_table_data(table) -> List[List[str]]:
    """提取 python-docx 表格对象为二维文本列表"""
    data = []
    for row in table.rows:
        row_data = [cell.text.strip() for cell in row.cells]
        data.append(row_data)
    return data


def is_heading_style(style_name: str) -> bool:
    """判断样式名是否为标题样式（如 'Heading 1'）"""
    return style_name.startswith("Heading")


def get_heading_level(style_name: str) -> int:
    """从 'Heading 2' 中提取级别数字，提取失败返回 0"""
    match = re.search(r"\d+", style_name)
    return int(match.group()) if match else 0


def parse_word_by_headings(doc_path: str) -> ChapterNode:
    """
    解析 Word 文档，按标题分章节提取文字和表格。
    返回一个根节点（level=0），其 children 为各顶级章节。
    """
    doc = Document(doc_path)
    root = ChapterNode(level=0)          # 虚拟根节点，容纳所有顶级内容
    stack: List[ChapterNode] = [root]    # 章节栈，栈顶为当前接收内容的章节

    # 遍历文档 body 中所有块级元素（段落和表格）
    for element in doc.element.body:
        # 判断是否为段落
        if element.tag.endswith("}p"):
            # 通过 XML 获取对应的 Paragraph 对象
            para = doc.paragraphs[0]  # 临时占位，无法直接映射，需要另一种方式
            # 注：直接遍历 doc.paragraphs 无法保持与表格的顺序，
            # 这里改用 iter_block_items 辅助函数（见下文）
            pass

    # 由于 python-docx 不直接提供“按文档流遍历段落和表格”的 API，
    # 我们可以利用辅助函数 iter_block_items（来自官方文档示例）
    from docx.oxml.ns import qn

    def iter_block_items(parent):
        """
        生成文档中段落和表格的顺序流（来自 python-docx 官方示例）
        """
        from docx.oxml.ns import nsdecls
        from docx.oxml import parse_xml
        for child in parent.element.body:
            if child.tag == qn('w:p'):
                yield doc.paragraphs[0]  # 同样面临映射问题……