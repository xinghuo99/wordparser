from docx import Document
from docx.oxml.ns import qn  # 用于按命名空间查找 XML 元素


def get_paragraph_display_text(paragraph):
    """
    提取单个段落中“可见”的文本。
    优先提取 <w:t> 文本（域结果或普通文字）；
    如果没有，降级提取 <w:instrText>（域代码）。
    """
    # 1. 尝试收集所有 <w:t> 的文本
    t_texts = []
    for run in paragraph.runs:
        for t_elem in run._element.findall(qn('w:t')):
            if t_elem.text:
                t_texts.append(t_elem.text)

    if t_texts:
        return ''.join(t_texts)

    # 2. 如果没有任何 <w:t>，则提取域代码 <w:instrText>
    instr_texts = []
    for run in paragraph.runs:
        for instr_elem in run._element.findall(qn('w:instrText')):
            if instr_elem.text:
                instr_texts.append(instr_elem.text)

    return ''.join(instr_texts)


def get_cell_all_text(cell):
    """
    递归提取单元格内所有文本，包括嵌套表格中的内容。
    """
    parts = []

    # 1. 提取当前单元格所有段落的显示文本
    for para in cell.paragraphs:
        text = get_paragraph_display_text(para)
        if text:
            parts.append(text)

    # 2. 递归处理嵌套在单元格内部的表格
    for nested_table in cell.tables:
        for row in nested_table.rows:
            for nested_cell in row.cells:
                nested_text = get_cell_all_text(nested_cell)
                if nested_text:
                    parts.append(nested_text)

    # 用换行符连接各部分，保留大致排版
    return '\n'.join(parts)


def extract_all_header_text(doc):
    """
    提取文档中所有节、所有类型页眉中的全部文本（含表格）
    """
    results = []
    for section in doc.sections:
        # 根据节属性收集对应的页眉对象
        headers = [section.header]
        if section.different_first_page_header_footer:
            headers.append(section.first_page_header)
        if section.different_odd_and_even_pages_header_footer:
            headers.append(section.even_page_header)

        for header in headers:
            # 若当前节的页眉链接到前一节，则实际内容在前一节，这里只做提醒
            # 实际读取链接的源头需要手动回溯 section 链表，此处略
            if header.is_linked_to_previous:
                print("⚠️ 当前节页眉链接到前一节，当前 header 可能为空，请手动获取前一节的页眉。")
                # 如需自动追溯，可通过 section._sectPr 的 previous_section 关系去查

            # 提取页眉中的段落
            for para in header.paragraphs:
                text = get_paragraph_display_text(para)
                if text.strip():
                    results.append(('[段落]', text))

            # 提取页眉中的表格（以及嵌套表格）
            for table in header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        cell_text = get_cell_all_text(cell)
                        if cell_text.strip():
                            results.append(('[表格单元格]', cell_text))

    return results


# ========== 使用示例 ==========
if __name__ == '__main__':
    doc = Document('your_document.docx')  # 替换为你的文件路径
    all_texts = extract_all_header_text(doc)

    for location, content in all_texts:
        print(f"{location}: {content}")