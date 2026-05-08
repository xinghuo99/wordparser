import pytest
import os
import sys
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..', 'src'))

from docx import Document
from word_parser import WordParser, Section
from document_creator import create_test_document


TEST_DOC_PATH = 'test_document.docx'


@pytest.fixture(scope='module')
def test_document():
    create_test_document(TEST_DOC_PATH)
    yield TEST_DOC_PATH
    if os.path.exists(TEST_DOC_PATH):
        os.remove(TEST_DOC_PATH)


def test_extract_sections(test_document):
    parser = WordParser(test_document)
    sections = parser.extract_sections()
    
    assert len(sections) == 4
    
    assert sections[0].title == '第一章 概述'
    assert sections[0].level == 1
    assert '第一章的正文内容' in sections[0].content
    assert sections[0].tables == []
    
    assert sections[1].title == '第二章 技术方案'
    assert sections[1].level == 1
    assert '技术方案的设计' in sections[1].content
    
    assert sections[2].title == '第三章 修订记录'
    assert sections[2].level == 1
    
    assert sections[3].title == '第四章 附录'
    assert sections[3].level == 1


def test_section_hierarchy(test_document):
    parser = WordParser(test_document)
    sections = parser.extract_sections()
    
    chapter1 = sections[0]
    assert len(chapter1.children) == 2
    
    section1_1 = chapter1.children[0]
    assert section1_1.title == '1.1 项目背景'
    assert section1_1.level == 2
    assert len(section1_1.tables) == 1
    assert section1_1.tables[0][0] == ['序号', '项目名称', '状态']
    
    section1_2 = chapter1.children[1]
    assert section1_2.title == '1.2 项目目标'
    assert section1_2.level == 2
    
    chapter2 = sections[1]
    assert len(chapter2.children) == 2
    
    section2_1 = chapter2.children[0]
    assert section2_1.title == '2.1 技术选型'
    assert len(section2_1.children) == 2
    
    section2_1_1 = section2_1.children[0]
    assert section2_1_1.title == '2.1.1 前端技术'
    assert section2_1_1.level == 3
    assert len(section2_1_1.tables) == 1
    
    section2_1_2 = section2_1.children[1]
    assert section2_1_2.title == '2.1.2 后端技术'
    assert section2_1_2.level == 3


def test_find_section_by_name(test_document):
    parser = WordParser(test_document)
    parser.extract_sections()
    
    found = parser.find_section_by_name('第一章 概述')
    assert found is not None
    assert found.title == '第一章 概述'
    
    found = parser.find_section_by_name('1.1 项目背景')
    assert found is not None
    assert found.title == '1.1 项目背景'
    assert len(found.tables) == 1
    
    found = parser.find_section_by_name('2.1.1 前端技术')
    assert found is not None
    assert found.title == '2.1.1 前端技术'
    
    found = parser.find_section_by_name('不存在的章节')
    assert found is None


def test_find_section_with_children(test_document):
    parser = WordParser(test_document)
    parser.extract_sections()
    
    found = parser.find_section_by_name('2.1 技术选型')
    assert found is not None
    assert len(found.children) == 2
    
    child_titles = [child.title for child in found.children]
    assert '2.1.1 前端技术' in child_titles
    assert '2.1.2 后端技术' in child_titles


def test_get_section_content(test_document):
    parser = WordParser(test_document)
    parser.extract_sections()
    
    section = parser.find_section_by_name('第一章 概述')
    content = parser.get_section_content(section)
    
    assert content['title'] == '第一章 概述'
    assert content['level'] == 1
    assert '正文内容' in content['content']
    assert len(content['children']) == 2
    
    child_titles = [child['title'] for child in content['children']]
    assert '1.1 项目背景' in child_titles
    assert '1.2 项目目标' in child_titles


def test_table_extraction(test_document):
    parser = WordParser(test_document)
    parser.extract_sections()
    
    section = parser.find_section_by_name('1.1 项目背景')
    assert len(section.tables) == 1
    
    table = section.tables[0]
    assert len(table) == 3
    assert table[0] == ['序号', '项目名称', '状态']
    assert table[1] == ['1', '项目A', '进行中']
    assert table[2] == ['2', '项目B', '已完成']


def test_extract_headers(test_document):
    parser = WordParser(test_document)
    headers = parser.extract_headers()
    
    assert len(headers) >= 1
    
    section_header = headers[0]
    assert 'section_index' in section_header
    assert 'headers' in section_header
    
    assert len(section_header['headers']) >= 1
    
    header = section_header['headers'][0]
    assert 'type' in header
    assert 'content' in header
    assert 'tables' in header
    assert header['type'] == 'default'
    assert 'Word Parser 测试文档' in header['content']


def test_get_all_headers_text(test_document):
    parser = WordParser(test_document)
    all_text = parser.get_all_headers_text()
    
    assert isinstance(all_text, str)
    assert 'Word Parser 测试文档' in all_text


def test_table_to_text(test_document):
    parser = WordParser(test_document)
    
    table = [
        ['序号', '项目名称', '状态'],
        ['1', '项目A', '进行中'],
        ['2', '项目B', '已完成']
    ]
    
    result = parser._table_to_text(table)
    
    assert isinstance(result, str)
    assert '序号: 1' in result
    assert '项目名称: 项目A' in result
    assert '状态: 进行中' in result
    assert '序号: 2' in result
    assert '项目名称: 项目B' in result
    assert '状态: 已完成' in result


def test_get_full_document_text(test_document):
    parser = WordParser(test_document)
    full_text = parser.get_full_document_text(include_headers=True)
    
    assert isinstance(full_text, str)
    assert '第一章 概述' in full_text
    assert '第二章 技术方案' in full_text
    assert '项目背景' in full_text
    assert '序号: 1' in full_text
    assert '项目名称: 项目A' in full_text
    assert 'Word Parser 测试文档' in full_text
    
    full_text_no_header = parser.get_full_document_text(include_headers=False)
    assert 'Word Parser 测试文档' not in full_text_no_header


def test_get_section_text_by_name(test_document):
    parser = WordParser(test_document)
    
    section_text = parser.get_section_text_by_name('第一章 概述')
    assert isinstance(section_text, str)
    assert '第一章 概述' in section_text
    assert '1.1 项目背景' in section_text
    assert '1.2 项目目标' in section_text
    assert '序号: 1' in section_text
    
    section_text_no_child = parser.get_section_text_by_name('第一章 概述', include_children=False)
    assert '第一章 概述' in section_text_no_child
    assert '1.1 项目背景' not in section_text_no_child
    
    section_text_child = parser.get_section_text_by_name('1.1 项目背景')
    assert '1.1 项目背景' in section_text_child
    assert '序号: 1' in section_text_child
    assert '项目名称: 项目A' in section_text_child
    
    not_found_text = parser.get_section_text_by_name('不存在的章节')
    assert '未找到章节' in not_found_text


def test_get_revision_records(test_document):
    parser = WordParser(test_document)
    
    revision_text, has_table = parser.get_revision_records()
    
    assert isinstance(revision_text, str)
    assert isinstance(has_table, bool)
    assert has_table == True
    assert '第三章 修订记录' in revision_text
    assert '本文档的修订历史记录如下' in revision_text
    assert '版本号: V1.0' in revision_text
    assert '修订日期: 2024-01-15' in revision_text
    assert '修订人: 张三' in revision_text
    assert '修订内容: 初始版本' in revision_text
    assert '版本号: V1.1' in revision_text
    assert '修订人: 李四' in revision_text


def test_get_revision_records_no_revision_section(tmp_path):
    """测试文档中没有修订记录章节的情况"""
    no_revision_doc = tmp_path / 'no_revision.docx'
    
    doc = Document()
    doc.add_heading('第一章 概述', level=1)
    doc.add_paragraph('这是测试文档')
    doc.save(str(no_revision_doc))
    
    parser = WordParser(str(no_revision_doc))
    revision_text, has_table = parser.get_revision_records()
    
    assert isinstance(revision_text, str)
    assert isinstance(has_table, bool)
    assert revision_text == ""
    assert has_table == False


def test_get_revision_records_as_subsection(tmp_path):
    """测试修订记录作为子章节的情况"""
    subsection_doc = tmp_path / 'subsection_revision.docx'
    
    doc = Document()
    doc.add_heading('附录', level=1)
    doc.add_heading('附录A 修订记录', level=2)
    doc.add_paragraph('这是子章节中的修订记录')
    
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = '版本'
    table.rows[0].cells[1].text = '内容'
    table.rows[1].cells[0].text = 'V1.0'
    table.rows[1].cells[1].text = '初始版本'
    
    doc.save(str(subsection_doc))
    
    parser = WordParser(str(subsection_doc))
    revision_text, has_table = parser.get_revision_records()
    
    assert isinstance(revision_text, str)
    assert isinstance(has_table, bool)
    assert has_table == True
    assert '附录A 修订记录' in revision_text
    assert '这是子章节中的修订记录' in revision_text
    assert '版本: V1.0' in revision_text


def test_get_revision_records_no_table(tmp_path):
    """测试修订记录没有表格的情况"""
    no_table_doc = tmp_path / 'no_table_revision.docx'
    
    doc = Document()
    doc.add_heading('修订记录', level=1)
    doc.add_paragraph('修订历史：')
    doc.add_paragraph('- V1.0: 初始版本')
    doc.add_paragraph('- V1.1: 更新内容')
    
    doc.save(str(no_table_doc))
    
    parser = WordParser(str(no_table_doc))
    revision_text, has_table = parser.get_revision_records()
    
    assert isinstance(revision_text, str)
    assert isinstance(has_table, bool)
    assert has_table == False
    assert '修订记录' in revision_text
    assert '修订历史：' in revision_text
    assert 'V1.0: 初始版本' in revision_text
    assert 'V1.1: 更新内容' in revision_text


def test_get_revision_records_as_paragraph_with_table(tmp_path):
    """测试修订记录是普通文字段落且后面紧跟表格的情况"""
    para_doc = tmp_path / 'paragraph_revision_with_table.docx'
    
    doc = Document()
    doc.add_heading('第一章 概述', level=1)
    doc.add_paragraph('这是第一章内容')
    doc.add_paragraph('修订记录')
    doc.add_paragraph('本文档的修订历史如下：')
    
    table = doc.add_table(rows=3, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '版本号'
    hdr_cells[1].text = '修订日期'
    hdr_cells[2].text = '修订人'
    hdr_cells[3].text = '修订内容'
    
    row_cells = table.rows[1].cells
    row_cells[0].text = 'V1.0'
    row_cells[1].text = '2024-01-15'
    row_cells[2].text = '张三'
    row_cells[3].text = '初始版本'
    
    row_cells = table.rows[2].cells
    row_cells[0].text = 'V1.1'
    row_cells[1].text = '2024-02-20'
    row_cells[2].text = '李四'
    row_cells[3].text = '更新内容'
    
    doc.add_heading('第二章 内容', level=1)
    doc.add_paragraph('这是第二章内容')
    
    doc.save(str(para_doc))
    
    parser = WordParser(str(para_doc))
    revision_text, has_table = parser.get_revision_records()
    
    assert isinstance(revision_text, str)
    assert isinstance(has_table, bool)
    assert has_table == True
    assert '修订记录' in revision_text
    assert '本文档的修订历史如下：' in revision_text
    assert '版本号: V1.0' in revision_text
    assert '修订日期: 2024-01-15' in revision_text
    assert '修订人: 张三' in revision_text
    assert '修订内容: 初始版本' in revision_text
    assert '版本号: V1.1' in revision_text
    assert '第二章 内容' not in revision_text


def test_get_revision_records_as_paragraph_without_table(tmp_path):
    """测试修订记录是普通文字段落且后面没有表格的情况"""
    para_no_table_doc = tmp_path / 'paragraph_revision_no_table.docx'
    
    doc = Document()
    doc.add_heading('第一章 概述', level=1)
    doc.add_paragraph('这是第一章内容')
    doc.add_paragraph('修订记录')
    doc.add_paragraph('- V1.0: 初始版本')
    doc.add_paragraph('- V1.1: 更新内容')
    doc.add_heading('第二章 内容', level=1)
    doc.add_paragraph('这是第二章内容')
    
    doc.save(str(para_no_table_doc))
    
    parser = WordParser(str(para_no_table_doc))
    revision_text, has_table = parser.get_revision_records()
    
    assert isinstance(revision_text, str)
    assert isinstance(has_table, bool)
    assert has_table == False
    assert '修订记录' in revision_text
    assert 'V1.0: 初始版本' in revision_text
    assert 'V1.1: 更新内容' in revision_text
    assert '第二章 内容' not in revision_text


def test_get_revision_records_paragraph_mixed_content(tmp_path):
    """测试修订记录是普通文字段落，后面跟着文字和表格的情况"""
    mixed_doc = tmp_path / 'paragraph_revision_mixed.docx'
    
    doc = Document()
    doc.add_heading('文档标题', level=1)
    doc.add_paragraph('文档内容')
    doc.add_paragraph('修订记录')
    doc.add_paragraph('以下是文档的修订历史：')
    doc.add_paragraph('重要说明：所有修订都经过审核。')
    
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = '版本'
    table.rows[0].cells[1].text = '描述'
    table.rows[1].cells[0].text = 'V1.0'
    table.rows[1].cells[1].text = '首次发布'
    
    doc.add_heading('附录', level=1)
    doc.add_paragraph('附录内容')
    
    doc.save(str(mixed_doc))
    
    parser = WordParser(str(mixed_doc))
    revision_text, has_table = parser.get_revision_records()
    
    assert isinstance(revision_text, str)
    assert isinstance(has_table, bool)
    assert has_table == True
    assert '修订记录' in revision_text
    assert '以下是文档的修订历史：' in revision_text
    assert '重要说明：所有修订都经过审核。' in revision_text
    assert '版本: V1.0' in revision_text
    assert '描述: 首次发布' in revision_text
    assert '附录' not in revision_text


def test_get_revision_records_prefer_section_over_paragraph(tmp_path):
    """测试文档同时包含章节类型和普通文字类型的修订记录时，优先提取章节类型"""
    mixed_revision_doc = tmp_path / 'mixed_revision_types.docx'
    
    doc = Document()
    doc.add_heading('第一章 概述', level=1)
    doc.add_paragraph('内容')
    
    doc.add_heading('第二章 修订记录', level=1)
    doc.add_paragraph('章节类型修订记录内容')
    
    section_table = doc.add_table(rows=2, cols=2)
    section_table.rows[0].cells[0].text = '版本'
    section_table.rows[0].cells[1].text = '内容'
    section_table.rows[1].cells[0].text = 'V1.0'
    section_table.rows[1].cells[1].text = '章节类型'
    
    doc.add_heading('第三章 附录', level=1)
    doc.add_paragraph('附录内容')
    doc.add_paragraph('修订记录')
    doc.add_paragraph('普通文字类型修订记录内容')
    
    para_table = doc.add_table(rows=2, cols=2)
    para_table.rows[0].cells[0].text = '版本'
    para_table.rows[0].cells[1].text = '内容'
    para_table.rows[1].cells[0].text = 'V2.0'
    para_table.rows[1].cells[1].text = '普通文字类型'
    
    doc.save(str(mixed_revision_doc))
    
    parser = WordParser(str(mixed_revision_doc))
    revision_text, has_table = parser.get_revision_records()
    
    assert isinstance(revision_text, str)
    assert isinstance(has_table, bool)
    assert has_table == True
    assert '第二章 修订记录' in revision_text
    assert '章节类型修订记录内容' in revision_text
    assert '版本: V1.0' in revision_text
    assert '内容: 章节类型' in revision_text
    assert 'V2.0' not in revision_text
    assert '普通文字类型' not in revision_text


def test_extract_sections_to_list(test_document):
    """测试按章节提取内容到列表"""
    parser = WordParser(test_document)
    sections_list = parser.extract_sections_to_list()
    
    assert isinstance(sections_list, list)
    assert len(sections_list) == 10
    
    for item in sections_list:
        assert isinstance(item, dict)
        assert 'title' in item
        assert 'content' in item
        assert isinstance(item['title'], str)
        assert isinstance(item['content'], str)
    
    chapter1 = sections_list[0]
    assert chapter1['title'] == '第一章 概述'
    assert '这是第一章的正文内容' in chapter1['content']
    assert '1.1 项目背景' in chapter1['content']
    assert '项目名称: 项目A' in chapter1['content']
    assert '1.2 项目目标' in chapter1['content']
    
    section_1_1 = sections_list[1]
    assert section_1_1['title'] == '1.1 项目背景'
    assert '本节介绍项目的背景信息' in section_1_1['content']
    assert '项目名称: 项目A' in section_1_1['content']
    
    chapter2 = sections_list[3]
    assert chapter2['title'] == '第二章 技术方案'
    assert '2.1.1 前端技术' in chapter2['content']
    assert '技术: React' in chapter2['content']
    
    revision = sections_list[8]
    assert revision['title'] == '第三章 修订记录'
    
    appendix = sections_list[9]
    assert appendix['title'] == '第四章 附录'
    assert '版本号: V1.0' in revision['content']
    assert '修订人: 张三' in revision['content']


def test_find_section_content_normal_case(test_document):
    """测试正常查找存在的章节"""
    parser = WordParser(test_document)
    sections_list = parser.extract_sections_to_list()
    
    content = WordParser.find_section_content(sections_list, '第一章 概述')
    assert isinstance(content, str)
    assert len(content) > 0
    assert '这是第一章的正文内容' in content
    assert '1.1 项目背景' in content
    
    content = WordParser.find_section_content(sections_list, '1.1 项目背景')
    assert isinstance(content, str)
    assert len(content) > 0
    assert '本节介绍项目的背景信息' in content
    assert '项目名称: 项目A' in content
    
    content = WordParser.find_section_content(sections_list, '第三章 修订记录')
    assert isinstance(content, str)
    assert len(content) > 0
    assert '版本号: V1.0' in content
    assert '修订人: 张三' in content


def test_find_section_content_not_found(test_document):
    """测试查找不存在的章节"""
    parser = WordParser(test_document)
    sections_list = parser.extract_sections_to_list()
    
    content = WordParser.find_section_content(sections_list, '不存在的章节')
    assert content == ""
    
    content = WordParser.find_section_content(sections_list, '第五章 参考文献')
    assert content == ""
    
    content = WordParser.find_section_content(sections_list, '')
    assert content == ""


def test_find_section_content_empty_list():
    """测试空列表输入"""
    content = WordParser.find_section_content([], '第一章 概述')
    assert content == ""
    
    content = WordParser.find_section_content(None, '第一章 概述')
    assert content == ""


def test_find_section_content_invalid_section_name():
    """测试无效的章节名称参数"""
    sections_list = [
        {'title': '第一章', 'content': '内容1'},
        {'title': '第二章', 'content': '内容2'}
    ]
    
    content = WordParser.find_section_content(sections_list, None)
    assert content == ""
    
    content = WordParser.find_section_content(sections_list, 123)
    assert content == ""
    
    content = WordParser.find_section_content(sections_list, "")
    assert content == ""


def test_find_section_content_invalid_list_format():
    """测试格式不正确的章节列表"""
    invalid_list = [
        {'title': '第一章', 'content': '内容1'},
        {'name': '第二章', 'content': '内容2'},
        {'title': '第三章'},
        {'content': '内容4'},
        '不是字典',
        None
    ]
    
    content = WordParser.find_section_content(invalid_list, '第一章')
    assert content == '内容1'
    
    content = WordParser.find_section_content(invalid_list, '第二章')
    assert content == ""
    
    content = WordParser.find_section_content(invalid_list, '第三章')
    assert content == ""
    
    content = WordParser.find_section_content(invalid_list, '第四章')
    assert content == ""


def test_find_section_content_empty_content():
    """测试章节内容为空的情况"""
    sections_list = [
        {'title': '空内容章节', 'content': ''},
        {'title': '有内容章节', 'content': '这是内容'}
    ]
    
    content = WordParser.find_section_content(sections_list, '空内容章节')
    assert content == ""
    
    content = WordParser.find_section_content(sections_list, '有内容章节')
    assert content == '这是内容'
